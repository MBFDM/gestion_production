import streamlit as st
import pandas as pd
import io
import tempfile
import os
import subprocess
import shutil
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from datetime import datetime
import locale
import sqlite3
import bcrypt
from functools import wraps
import base64
import time
import platform

# ------------------ Configuration initiale de la page ------------------
st.set_page_config(
    page_title="Générateur de certificats - Authentification", 
    layout="wide",
    page_icon="logo_1.jpg",
    initial_sidebar_state="expanded"
)

# ------------------ Logo en base64 ------------------
def get_logo_base64():
    """Crée un logo SVG simple en base64"""
    svg_logo = '''
    <svg xmlns="http://www.w3.org/2000/svg" viewBox="0 0 200 80" width="200" height="80">
        <rect x="0" y="0" width="200" height="80" rx="10" fill="#1e3a5f"/>
        <text x="100" y="35" font-family="Arial, sans-serif" font-size="24" font-weight="bold" fill="#ffffff" text-anchor="middle">
            📄 CertiGen
        </text>
        <text x="100" y="60" font-family="Arial, sans-serif" font-size="14" fill="#88b4d8" text-anchor="middle">
            Générateur de certificats
        </text>
    </svg>
    '''
    return base64.b64encode(svg_logo.encode()).decode()

# ------------------ Métriques de performance ------------------
class PerformanceMetrics:
    def __init__(self):
        self.start_time = None
        self.certificates_generated = 0
        self.processing_times = []
        self.session_start = time.time()
        
    def start_processing(self):
        self.start_time = time.time()
        
    def end_processing(self, count):
        self.certificates_generated += count
        elapsed = time.time() - self.start_time
        self.processing_times.append(elapsed)
        return elapsed
    
    def get_metrics(self):
        total_time = time.time() - self.session_start
        avg_time = sum(self.processing_times) / len(self.processing_times) if self.processing_times else 0
        return {
            "Total générés": self.certificates_generated,
            "Temps moyen": f"{avg_time:.2f}s" if self.processing_times else "0s",
            "Dernier temps": f"{self.processing_times[-1]:.2f}s" if self.processing_times else "0s",
            "Total sessions": len(self.processing_times),
            "Temps total session": f"{total_time:.1f}s"
        }
    
    def get_system_info(self):
        return {
            "OS": platform.system(),
            "Python": platform.python_version(),
            "Architecture": platform.machine()
        }

# Initialiser les métriques
if 'metrics' not in st.session_state:
    st.session_state.metrics = PerformanceMetrics()

# ------------------ Gestion de la base de données SQLite ------------------
DB_PATH = "users.db"

def init_db():
    """Crée la table users et le compte admin par défaut si inexistant."""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    # Vérifier si un admin existe déjà
    c.execute("SELECT * FROM users WHERE role = 'admin'")
    if not c.fetchone():
        # Créer admin par défaut : username = admin, password = admin123
        hashed = bcrypt.hashpw("admin123".encode(), bcrypt.gensalt()).decode()
        c.execute("INSERT INTO users (username, password_hash, role) VALUES (?, ?, ?)",
                  ("admin", hashed, "admin"))
        conn.commit()
    conn.close()

def get_db_connection():
    return sqlite3.connect(DB_PATH)

def verify_password(username, password):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT password_hash, role FROM users WHERE username = ?", (username,))
    row = c.fetchone()
    conn.close()
    if row:
        stored_hash = row[0]
        role = row[1]
        if bcrypt.checkpw(password.encode(), stored_hash.encode()):
            return role
    return None

def add_user(username, password, role):
    hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    conn = get_db_connection()
    c = conn.cursor()
    try:
        c.execute("INSERT INTO users (username, password_hash, role) VALUES (?, ?, ?)",
                  (username, hashed, role))
        conn.commit()
        return True
    except sqlite3.IntegrityError:
        return False
    finally:
        conn.close()

def update_user_password(username, new_password):
    hashed = bcrypt.hashpw(new_password.encode(), bcrypt.gensalt()).decode()
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("UPDATE users SET password_hash = ? WHERE username = ?", (hashed, username))
    conn.commit()
    conn.close()

def delete_user(username):
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("DELETE FROM users WHERE username = ? AND role != 'admin'", (username,))
    conn.commit()
    conn.close()

def get_all_users():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT id, username, role, created_at FROM users ORDER BY id")
    users = c.fetchall()
    conn.close()
    return users

def get_user_count():
    conn = get_db_connection()
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM users")
    count = c.fetchone()[0]
    conn.close()
    return count

# ------------------ Fonctions de l'application principale (inchangées) ------------------
champs_cotes = [
    "N° Assuré", "N° Police", "N° Référence", "Intermédiaire", "Tél", "Tél WhatApps",
    "Nom(s) et Prénoms", "Date de Naissance", "Sexe", "Effet", "Echéance", "Durée (mois)",
    "Fractionnement", "Date de souscription", "Périodicité"
]

champs_dessous = [
    "Garantie", "Capital (FCFA)", "Primes Périodes (FCFA)",
    "Prime nette", "Accessoires", "Prime Totale"
]

champs_attendus = champs_cotes + champs_dessous
champs_date = ["Date de Naissance", "Effet", "Echéance", "Date de souscription"]
champs_decalage_double = ["N° Référence", "Nom(s) et Prénoms", "Date de Naissance"]
champs_decalage_triple = ["Date de souscription"]

def formater_date(valeur):
    if pd.isna(valeur):
        return ""
    if isinstance(valeur, (pd.Timestamp, datetime)):
        try:
            locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
        except:
            try:
                locale.setlocale(locale.LC_TIME, 'fr_FR')
            except:
                pass
        return valeur.strftime("%d %B %Y")
    if isinstance(valeur, str):
        formats = ["%Y-%m-%d", "%d/%m/%Y", "%m/%d/%Y", "%d-%m-%Y", "%Y/%m/%d", "%d.%m.%Y"]
        for fmt in formats:
            try:
                dt = datetime.strptime(valeur, fmt)
                try:
                    locale.setlocale(locale.LC_TIME, 'fr_FR.UTF-8')
                except:
                    try:
                        locale.setlocale(locale.LC_TIME, 'fr_FR')
                    except:
                        pass
                return dt.strftime("%d %B %Y")
            except:
                continue
        return valeur
    return str(valeur)

def convert_docx_to_pdf(docx_path, pdf_path):
    libreoffice_cmds = ['libreoffice', 'soffice']
    for cmd in libreoffice_cmds:
        if shutil.which(cmd):
            try:
                subprocess.run(
                    [cmd, '--headless', '--convert-to', 'pdf', '--outdir', os.path.dirname(pdf_path), docx_path],
                    check=True, capture_output=True
                )
                generated_pdf = os.path.join(os.path.dirname(pdf_path),
                                             os.path.splitext(os.path.basename(docx_path))[0] + '.pdf')
                if os.path.exists(generated_pdf):
                    os.rename(generated_pdf, pdf_path)
                return True
            except Exception:
                continue
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        return True
    except Exception:
        return False

def remplir_un_certificat(template_bytes, data_dict, style_config):
    template_stream = BytesIO(template_bytes)
    doc = Document(template_stream)

    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell_text = cell.text.strip()
                for champ, valeur in data_dict.items():
                    if champ in cell_text or cell_text == champ:
                        if champ in champs_cotes:
                            target_row = row_idx
                            if champ in champs_decalage_triple:
                                target_col = col_idx + 3
                            elif champ in champs_decalage_double:
                                target_col = col_idx + 2
                            else:
                                target_col = col_idx + 1
                            while target_col >= len(table.rows[target_row].cells):
                                for r in table.rows:
                                    r.cells.add()
                            target_cell = table.rows[target_row].cells[target_col]
                        elif champ in champs_dessous:
                            target_row = row_idx + 1
                            target_col = col_idx
                            while target_row >= len(table.rows):
                                table.add_row()
                            target_cell = table.rows[target_row].cells[target_col]
                        else:
                            continue
                        target_cell.text = ""
                        paragraph = target_cell.paragraphs[0]
                        if style_config['alignment'] == 'gauche':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        elif style_config['alignment'] == 'centre':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        elif style_config['alignment'] == 'droite':
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        run = paragraph.add_run(str(valeur))
                        if style_config['font_name']:
                            run.font.name = style_config['font_name']
                        if style_config['font_size']:
                            run.font.size = Pt(style_config['font_size'])
                        if style_config['color_hex']:
                            rgb = RGBColor(
                                int(style_config['color_hex'][1:3], 16),
                                int(style_config['color_hex'][3:5], 16),
                                int(style_config['color_hex'][5:7], 16)
                            )
                            run.font.color.rgb = rgb
                        run.font.bold = style_config['bold']
                        run.font.italic = style_config['italic']
                        break

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

def generer_tous_certificats(template_bytes, df, style_config):
    certificats = []
    for idx, row in df.iterrows():
        data_dict = {}
        for champ in champs_attendus:
            valeur = row[champ]
            if pd.notna(valeur):
                if champ in champs_date:
                    valeur_formatee = formater_date(valeur)
                else:
                    valeur_formatee = str(valeur)
            else:
                valeur_formatee = ""
            data_dict[champ] = valeur_formatee
        identifiant = data_dict.get("Nom(s) et Prénoms", f"ligne_{idx+1}").replace("/", "_")
        docx_bytesio = remplir_un_certificat(template_bytes, data_dict, style_config)
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(docx_bytesio.getvalue())
            tmp_docx_path = tmp_docx.name
        pdf_path = tmp_docx_path.replace(".docx", ".pdf")
        conversion_ok = convert_docx_to_pdf(tmp_docx_path, pdf_path)
        pdf_bytes = None
        if conversion_ok and os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            os.unlink(pdf_path)
        os.unlink(tmp_docx_path)
        certificats.append((idx, identifiant, docx_bytesio, pdf_bytes))
    return certificats

def page_generateur():
    st.title("📄 Générateur de certificats (Word + PDF) personnalisables")
    
    # Afficher le logo dans la sidebar
    with st.sidebar:
        logo_base64 = get_logo_base64()
        st.markdown(f"""
        <div style="text-align: center; padding: 10px;">
            <img src="data:image/svg+xml;base64,{logo_base64}" style="width: 100%; max-width: 200px;">
        </div>
        """, unsafe_allow_html=True)
        st.markdown("---")
        
        # Métriques système
        st.subheader("📊 Métriques système")
        sys_info = st.session_state.metrics.get_system_info()
        st.metric("💻 Système", sys_info["OS"])
        st.metric("🐍 Python", sys_info["Python"])
        st.metric("🔧 Architecture", sys_info["Architecture"])
        
        st.markdown("---")
        
        # Métriques de l'application
        st.subheader("📈 Métriques application")
        metrics = st.session_state.metrics.get_metrics()
        
        # Métriques en colonnes
        col1, col2 = st.columns(2)
        with col1:
            st.metric("📄 Certificats générés", metrics["Total générés"], delta=None)
            st.metric("⏱️ Dernier temps", metrics["Dernier temps"])
        with col2:
            st.metric("🔄 Sessions", metrics["Total sessions"])
            st.metric("⚡ Temps moyen", metrics["Temps moyen"])
        
        st.metric("⏱️ Temps total session", metrics["Temps total session"])
    
    st.markdown("""
    Chargez un modèle Word (avec tableaux contenant les libellés) et un fichier Excel.
    - Les champs de la liste **côté** (`champs_cotes`) sont insérés **à droite** du libellé.
    - Les champs `N° Référence`, `Nom(s) et Prénoms` et `Date de souscription` sont insérés **deux/trois cellules à droite**.
    - Les champs de la liste **dessous** (`champs_dessous`) sont insérés **en dessous** du libellé.
    - Les dates sont formatées en **JJ Mois AAAA** (ex: 12 Mars 2000).
    """)

    col1, col2 = st.columns(2)
    with col1:
        modele_file = st.file_uploader("📄 Modèle Word (.docx)", type=["docx"])
    with col2:
        excel_file = st.file_uploader("📊 Fichier Excel (.xlsx)", type=["xlsx"])

    # Barre latérale pour la personnalisation
    st.sidebar.header("🎨 Personnalisation des valeurs insérées")
    font_name = st.sidebar.selectbox("Police", ["Arial", "Times New Roman", "Calibri", "Verdana", "Courier New"], index=0)
    font_size = st.sidebar.slider("Taille (pt)", 8, 48, 11)
    color_hex = st.sidebar.color_picker("Couleur du texte", "#000000")
    bold = st.sidebar.checkbox("Gras", value=False)
    italic = st.sidebar.checkbox("Italique", value=False)
    alignment = st.sidebar.radio("Alignement horizontal", ["gauche", "centre", "droite"], index=0)

    style_config = {
        'font_name': font_name,
        'font_size': font_size,
        'color_hex': color_hex,
        'bold': bold,
        'italic': italic,
        'alignment': alignment
    }

    if modele_file and excel_file:
        try:
            df = pd.read_excel(excel_file, engine='openpyxl')
            st.success(f"Excel chargé : {df.shape[0]} ligne(s), {df.shape[1]} colonne(s)")
            
            # Afficher les stats du fichier
            col_stats1, col_stats2, col_stats3 = st.columns(3)
            with col_stats1:
                st.metric("📊 Lignes", df.shape[0])
            with col_stats2:
                st.metric("📋 Colonnes", df.shape[1])
            with col_stats3:
                non_empty = df.count().sum()
                st.metric("✅ Cellules remplies", non_empty)
            
            st.subheader("Aperçu du fichier Excel")
            st.dataframe(df, use_container_width=True)

            colonnes_manquantes = [champ for champ in champs_attendus if champ not in df.columns]
            if colonnes_manquantes:
                st.error(f"❌ Colonnes manquantes : {', '.join(colonnes_manquantes)}")
                st.stop()
            else:
                st.success("✅ Tous les en-têtes requis sont présents.")

            # Démarrer le chrono
            st.session_state.metrics.start_processing()
            
            with st.spinner(f"Génération de {df.shape[0]} certificat(s)..."):
                template_bytes = modele_file.read()
                certificats = generer_tous_certificats(template_bytes, df, style_config)
            
            # Enregistrer les métriques
            elapsed = st.session_state.metrics.end_processing(len(certificats))
            
            st.success(f"✅ {len(certificats)} certificat(s) généré(s) en {elapsed:.2f} secondes.")

            if len(certificats) > 0:
                first_docx = certificats[0][2]
                st.download_button(
                    label="📄 Télécharger le Modèle Word Final (exemple première ligne)",
                    data=first_docx.getvalue(),
                    file_name="modele_word_final.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            st.subheader("📑 Certificats générés")
            zip_word = BytesIO()
            zip_pdf = BytesIO()
            with zipfile.ZipFile(zip_word, 'w') as zw:
                with zipfile.ZipFile(zip_pdf, 'w') as zp:
                    for idx, ident, docx_bytesio, pdf_bytes in certificats:
                        safe_name = ident.replace(" ", "_").replace("(", "").replace(")", "")
                        docx_name = f"{safe_name}.docx"
                        pdf_name = f"{safe_name}.pdf"
                        zw.writestr(docx_name, docx_bytesio.getvalue())
                        if pdf_bytes:
                            zp.writestr(pdf_name, pdf_bytes)
                        col_a, col_b, col_c, col_d = st.columns([3,1,1,1])
                        col_a.write(f"**{ident}**")
                        col_b.download_button("📄 Word", data=docx_bytesio.getvalue(), file_name=docx_name, key=f"word_{idx}")
                        if pdf_bytes:
                            col_c.download_button("📑 PDF", data=pdf_bytes, file_name=pdf_name, key=f"pdf_{idx}")
                        else:
                            col_c.write("❌ PDF non généré")
            zip_word.seek(0)
            zip_pdf.seek(0)
            st.markdown("---")
            col_zip1, col_zip2 = st.columns(2)
            with col_zip1:
                st.download_button("📦 Tous les Word (ZIP)", data=zip_word, file_name="tous_word.zip", mime="application/zip")
            with col_zip2:
                st.download_button("📦 Tous les PDF (ZIP)", data=zip_pdf, file_name="tous_pdf.zip", mime="application/zip", disabled=(zip_pdf.getbuffer().nbytes == 0))

        except Exception as e:
            st.error(f"❌ Erreur : {str(e)}")

def page_admin():
    st.title("👑 Administration des utilisateurs")
    st.markdown("Gérer les comptes utilisateurs de l'application.")
    
    # Métriques admin
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("👥 Utilisateurs totaux", get_user_count())
    with col2:
        st.metric("📄 Certificats générés", st.session_state.metrics.certificates_generated)
    with col3:
        st.metric("🔄 Sessions", st.session_state.metrics.get_metrics()["Total sessions"])

    # Ajouter un utilisateur
    with st.expander("➕ Ajouter un utilisateur"):
        new_username = st.text_input("Nom d'utilisateur")
        new_password = st.text_input("Mot de passe", type="password")
        new_role = st.selectbox("Rôle", ["user", "admin"])
        if st.button("Créer l'utilisateur"):
            if new_username and new_password:
                if add_user(new_username, new_password, new_role):
                    st.success(f"✅ Utilisateur '{new_username}' créé avec succès.")
                else:
                    st.error("❌ Nom d'utilisateur déjà existant.")
            else:
                st.warning("⚠️ Veuillez remplir tous les champs.")

    # Liste des utilisateurs
    users = get_all_users()
    st.subheader("📋 Liste des utilisateurs")
    if users:
        user_df = pd.DataFrame(users, columns=["ID", "Nom d'utilisateur", "Rôle", "Créé le"])
        st.dataframe(user_df, use_container_width=True)

        # Modification / suppression
        st.subheader("🔧 Modifier ou supprimer un utilisateur")
        # Filtrer pour ne pas afficher l'utilisateur actuel
        user_list = [u[1] for u in users if u[1] != st.session_state.username]
        if user_list:
            selected_username = st.selectbox("Choisir un utilisateur", user_list)
            if selected_username:
                col1, col2 = st.columns(2)
                with col1:
                    new_pass = st.text_input("Nouveau mot de passe (laisser vide pour ne pas changer)", type="password", key="admin_new_pass")
                    if st.button("🔄 Changer le mot de passe"):
                        if new_pass:
                            update_user_password(selected_username, new_pass)
                            st.success("✅ Mot de passe mis à jour.")
                        else:
                            st.info("ℹ️ Aucun changement.")
                with col2:
                    if st.button("🗑️ Supprimer cet utilisateur"):
                        if selected_username != "admin":
                            delete_user(selected_username)
                            st.success(f"✅ Utilisateur '{selected_username}' supprimé.")
                            st.rerun()
                        else:
                            st.error("❌ Impossible de supprimer le compte admin par défaut.")
        else:
            st.info("ℹ️ Aucun autre utilisateur à gérer.")
    else:
        st.info("ℹ️ Aucun utilisateur trouvé.")

# ------------------ Gestion de l'authentification ------------------
def login_page():
    # Afficher le logo sur la page de login
    logo_base64 = get_logo_base64()
    st.markdown(f"""
    <div style="text-align: center; padding: 20px;">
        <img src="data:image/svg+xml;base64,{logo_base64}" style="width: 300px;">
    </div>
    """, unsafe_allow_html=True)
    
    st.title("🔐 Connexion")
    
    # Métriques sur la page de login
    with st.expander("📊 Statistiques de l'application"):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("👥 Utilisateurs", get_user_count())
        with col2:
            st.metric("📄 Certificats générés", st.session_state.metrics.certificates_generated if hasattr(st.session_state, 'metrics') else 0)
        with col3:
            total_sessions = st.session_state.metrics.get_metrics()["Total sessions"] if hasattr(st.session_state, 'metrics') else 0
            st.metric("🔄 Sessions", total_sessions)
    
    username = st.text_input("👤 Nom d'utilisateur")
    password = st.text_input("🔑 Mot de passe", type="password")
    
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        if st.button("🔓 Se connecter", use_container_width=True):
            role = verify_password(username, password)
            if role:
                st.session_state.logged_in = True
                st.session_state.username = username
                st.session_state.role = role
                st.success(f"✅ Bienvenue {username} ({role})")
                time.sleep(0.5)
                st.rerun()
            else:
                st.error("❌ Nom d'utilisateur ou mot de passe incorrect")

def logout():
    for key in ['logged_in', 'username', 'role']:
        if key in st.session_state:
            del st.session_state[key]
    st.rerun()

# ------------------ Initialisation ------------------
init_db()

# État de connexion
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    login_page()
else:
    # Barre latérale avec déconnexion
    st.sidebar.markdown(f"**Connecté :** {st.session_state.username} ({st.session_state.role})")
    if st.sidebar.button("🚪 Déconnexion", use_container_width=True):
        logout()

    # Menu principal
    st.sidebar.title("Navigation")
    menu = ["Générateur de certificats"]
    if st.session_state.role == "admin":
        menu.append("Administration")
    choice = st.sidebar.radio("Aller à", menu, index=0)

    if choice == "Générateur de certificats":
        page_generateur()
    elif choice == "Administration":
        page_admin()
